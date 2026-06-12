from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup (A4, RTL) ───────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin   = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── RTL helpers ────────────────────────────────────────────────────────────

def set_rtl_para(paragraph):
    """Add <w:bidi/> to paragraph pPr and align right."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_rtl_run(run):
    """Add <w:rtl/> and <w:lang w:bidi='he-IL'/> to run rPr."""
    rPr = run._r.get_or_add_rPr()
    rtl_el = OxmlElement('w:rtl')
    rPr.append(rtl_el)
    lang_el = OxmlElement('w:lang')
    lang_el.set(qn('w:bidi'), 'he-IL')
    rPr.append(lang_el)

def set_rtl_cell(cell):
    """Set RTL on all paragraphs in a table cell."""
    for p in cell.paragraphs:
        set_rtl_para(p)
        for run in p.runs:
            set_rtl_run(run)

def set_table_rtl(table):
    """Set bidi on table and all cells."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    bidi = OxmlElement('w:bidiVisual')
    tbl_pr.append(bidi)
    for row in table.rows:
        for cell in row.cells:
            set_rtl_cell(cell)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    set_rtl_para(p)
    for run in p.runs:
        set_rtl_run(run)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, size=11, rtl=True, color=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    if rtl:
        set_rtl_para(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if rtl:
        set_rtl_run(run)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl_para(p)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    set_rtl_run(run)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 60)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(200, 200, 200)
    return p

# ═══════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run('praxisAI')
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(37, 99, 235)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('מסמך אפיון מערכת — גרסה 1.0')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(71, 85, 105)
set_rtl_run(r)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run('פלטפורמת תיעוד קליני מבוססת בינה מלאכותית לאנשי מקצוע פרא-רפואיים בישראל')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(100, 116, 139)
set_rtl_run(r)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run('יוני 2026')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(148, 163, 184)
set_rtl_run(r)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. מבוא ומטרות
# ═══════════════════════════════════════════════════════════════════════════
heading('1. מבוא ומטרות')

para('praxisAI היא פלטפורמת תיעוד קליני מבוססת בינה מלאכותית, המיועדת לקליניקות פרא-רפואיות בישראל: פיזיותרפיה, קלינאות תקשורת, ריפוי בעיסוק, תזונה קלינית, סיעוד, פסיכולוגיה קלינית וכירופרקטיקה. המטרה: לאפשר למטפל להתמקד בטיפול — בעוד ה-AI מטפל בתיעוד.')

heading('1.1 הבעיה שאנחנו פותרים', level=2)
bullet('אנשי מקצוע פרא-רפואיים מבזבזים 20–40% מזמנם על תיעוד מינהלי')
bullet('רשומות ידניות איכותיות דורשות 10–20 דקות לטיפול')
bullet('שגיאות ושכחה בתיעוד בדיעבד')
bullet('מסמכים לביטוח לאומי, הפניות ודו"חות — עבודה ידנית נוספת')

heading('1.2 הפתרון', level=2)
para('הקלט → תמלול בעברית ב-AI → רשומה קלינית מובנית → מסמכים אוטומטיים')
para('שחרור 5–10 שעות עבודה ידנית בשבוע לכל מטפל.')

divider()

# ═══════════════════════════════════════════════════════════════════════════
# 2. קהל יעד
# ═══════════════════════════════════════════════════════════════════════════
heading('2. קהל יעד')

heading('2.1 משתמשים עיקריים', level=2)
bullet('אנשי מקצוע פרא-רפואיים — מטפלים המשתמשים ב-Scribe ובתיעוד')
bullet('מנהלי קליניקה (בעלים / אדמין) — ניהול צוות, הגדרות, אנליטיקות')
bullet('פקידות קבלה — גישה מוגבלת לפרטי מטופלים')

heading('2.2 פרופיל קליניקה טיפוסי', level=2)
bullet('1–15 מטפלים')
bullet('50–500 מטופלים פעילים')
bullet('ישראל, טיפול בשפה העברית')
bullet('תחומים: פיזיותרפיה, קלינאות תקשורת, ריפוי בעיסוק, תזונה קלינית, סיעוד, פסיכולוגיה קלינית, כירופרקטיקה')

divider()

# ═══════════════════════════════════════════════════════════════════════════
# 3. מודולים ויכולות
# ═══════════════════════════════════════════════════════════════════════════
heading('3. מודולים ויכולות')

# 3.1 Auth
heading('3.1 אימות משתמשים', level=2)
bullet('כניסה עם גוגל (OAuth) — חובה אישור מראש')
bullet('מנגנון הזמנה-בלבד: רק משתמשים שהוזמנו במייל יכולים לגשת')
bullet('חסימת הרשמה חופשית — הגנה על נתוני מטופלים')
bullet('Super Admin: גישה לכל הקליניקות במערכת')

# 3.2 ניהול קליניקות
heading('3.2 ניהול קליניקות רב-אתריות', level=2)
bullet('קליניקה אחת או יותר תחת אותה חשבון')
bullet('החלפה מהירה בין קליניקות (Clinic Switcher בסרגל הצד)')
bullet('הגדרות נפרדות לכל קליניקה (תבנית תיעוד, לוגו)')
bullet('יצירת קליניקה חדשה ע"י Super Admin בלבד')
bullet('Super Admin: מחיקת קליניקה עם אישור שם מדויק')

# 3.3 ניהול משתמשים
heading('3.3 ניהול משתמשים והרשאות', level=2)
bullet('תפקידים: בעלים, מנהל, מטפל/ת, קבלה')
bullet('הזמנת חבר צוות במייל')
bullet('הפעלה / השהיה / מחיקה של חבר צוות')
bullet('הצגת "התחבר לאחרונה" — זמן יחסי בעברית')
bullet('חסימת מחיקה עצמית ומחיקת הבעלים')

# 3.4 ניהול מטופלים
heading('3.4 ניהול מטופלים', level=2)
bullet('יצירה ועדכון: שם, ת.ז., תאריך לידה, טלפון, מייל, קופת חולים, אבחנה')
bullet('סימון "תביעת ביטוח לאומי" לכל מטופל')
bullet('סטטוס: פעיל / שוחרר / בהמתנה')
bullet('היסטוריית טיפולים כרונולוגית עם תרשים VAS')
bullet('קישור ישיר למסמכים שנוצרו')

# 3.5 AI Scribe
heading('3.5 AI Scribe — תיעוד קולי', level=2)
para('הלב של המוצר. המטפל מקליט את שיחת הטיפול — ה-AI מתמלל וכותב רשומה קלינית מותאמת לתחום.', bold=True)

p = doc.add_paragraph()
set_rtl_para(p)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
r1 = p.add_run('זרימת עבודה: ')
r1.bold = True
r1.font.size = Pt(11)
set_rtl_run(r1)
r2 = p.add_run('הקלטה ← תמלול עברית (Deepgram Nova-2 / Whisper) ← יצירת רשומה (Claude AI) ← עריכה ← שמירה')
r2.font.size = Pt(11)
set_rtl_run(r2)

bullet('מצב פקודה קולית: "פרקסיס תרשום" — מתחיל הקלטה; "פרקסיס סיים" — עוצר ומייצר רשומה')
bullet('מצב ידני: לחיצת כפתור להפעלה ועצירה')
bullet('ויזואליזר גלים חי בזמן ההקלטה')
bullet('עריכת הרשומה לפני שמירה')
bullet('שמירת VAS (דרגת כאב 0–10) לכל טיפול')

# 3.6 Templates
heading('3.6 תבניות תיעוד קליניות', level=2)
para('הרשומה הקלינית מותאמת לסוג הקליניקה — Claude מקבל הנחיות מקצועיות ספציפיות לכל תחום.')

para('פיזיותרפיה:', bold=True, space_after=2)
bullet('🦴  אורטופדי אמבולטורי — ROM, MMT, Lachman, SLR ועוד')
bullet('🧠  נוירולוגי אמבולטורי — Barthel, FIM, Berg, MAS, MMSE')
bullet('🧒  ילדים והתפתחות — GMFCS, GMFM, PEDI, מיילסטונים')
bullet('🔄  שיקום — STG/LTG מדידים, FIM, Barthel, תכנון שחרור')
bullet('🏥  בית חולים — אקוטי — Weight-bearing, נוהלי בטיחות, Morse fall')
bullet('🌸  רצפת אגן — Oxford 0-5, PFDI-20, ביופידבק')
bullet('🌀  וסטיבולרי — Dix-Hallpike, Epley, DHI, nystagmus')

para('קלינאות תקשורת:', bold=True, space_before=6, space_after=2)
bullet('🗣️  מבוגרים — אפזיה, דיספגיה, קול, WAB-R, ASHA NOMS, DOSS')
bullet('🌱  ילדים — עיכוב שפה, ASD, פונולוגיה, PLS-5, GFTA-3, CELF')

para('ריפוי בעיסוק:', bold=True, space_before=6, space_after=2)
bullet('🤲  מבוגרים — ADL/IADL, כוח יד, קוגניציה, COPM, AMPS, FIM')
bullet('🎨  ילדים — עיבוד חושי SI, DCD, כתיבה, SPM, MABC-2, BOT-2')

para('תחומים נוספים:', bold=True, space_before=6, space_after=2)
bullet('🥗  תזונה קלינית — BMI, SGA, MNA, Harris-Benedict, תוכנית תפריט')
bullet('💊  סיעוד קהילתי — סימנים חיוניים, PUSH score, NANDA, Barthel/Katz')
bullet('🧩  פסיכולוגיה קלינית — MSE, PHQ-9/GAD-7, CBT/DBT פרמולציה, C-SSRS')
bullet('🔧  כירופרקטיקה — Subluxation complex, HVLA, Kemp, Spurling, ROM')
bullet('📋  תבנית מותאמת — הדבקת דגימה → Claude מנתח → מגדיר סעיפים אוטומטית')

para('שינוי תבנית חל מיידית: טופס תיעוד ידני, בקשת AI Scribe, ועזרת AI בצ\'אט.', bold=False)

# 3.7 Documents
heading('3.7 יצירת מסמכים אוטומטית', level=2)
bullet('מכתב לביטוח לאומי')
bullet('הפניה לרופא / ספציאליסט')
bullet('דו"ח התקדמות')
bullet('סיכום שחרור')
bullet('תביעת ביטוח')
bullet('אישור מחלה')
para('מסמכים נשמרים כטיוטה / סופי, מקושרים למטופל.')

# 3.8 Chat
heading('3.8 צ\'אט AI קליני', level=2)
bullet('עוזר AI בעברית המתמחה בתחום הפרא-רפואי')
bullet('עונה על שאלות פרוטוקול, אנטומיה, פתולוגיות')
bullet('מכיר את תבנית התיעוד של הקליניקה')
bullet('מוגבל לנושאים קליניים בלבד')

# 3.9 Analytics
heading('3.9 אנליטיקות', level=2)
bullet('תרשים מגמת VAS למטופל בודד (לפחות 2 נקודות נדרשות)')
bullet('סה"כ טיפולים לפי תקופה')
bullet('(מתוכנן) השוואת מטפלים, ימי שיא, ממוצע טיפולים למטופל')

# 3.10 Onboarding
heading('3.10 מרכז קליטה (Onboarding)', level=2)
bullet('סיור כניסה ראשון — 4 שקופיות המסבירות את המוצר')
bullet('מחוון התקדמות עגול (progress ring) בפינה שמאל-תחתית')
bullet('צ\'קליסט: מטופל ראשון, טיפול ראשון, מסמך ראשון, תבנית, הזמנת צוות')
bullet('מתעדכן אוטומטית עם כל פעולה שמתבצעת')
bullet('ניתן לסגירה לצמיתות לאחר השלמה')

divider()

# ═══════════════════════════════════════════════════════════════════════════
# 4. תפקידים והרשאות
# ═══════════════════════════════════════════════════════════════════════════
heading('4. תפקידים והרשאות')

table = doc.add_table(rows=6, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['יכולת', 'בעלים', 'מנהל', 'מטפל/ת', 'קבלה']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_run(cell.paragraphs[0].runs[0])

rows_data = [
    ['ניהול מטופלים',      '✓', '✓', '✓', '✓'],
    ['AI Scribe + תיעוד',  '✓', '✓', '✓', '✗'],
    ['יצירת מסמכים',       '✓', '✓', '✓', '✗'],
    ['ניהול משתמשים',      '✓', '✓', '✗', '✗'],
    ['הגדרות תבנית',       '✓', '✓', '✗', '✗'],
]
for r_idx, row_data in enumerate(rows_data):
    for c_idx, val in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell.paragraphs[0].runs:
            set_rtl_run(cell.paragraphs[0].runs[0])

set_table_rtl(table)
doc.add_paragraph()

divider()

# ═══════════════════════════════════════════════════════════════════════════
# 5. ארכיטקטורה טכנית
# ═══════════════════════════════════════════════════════════════════════════
heading('5. ארכיטקטורה טכנית')

heading('5.1 Stack', level=2)
bullet('Frontend: Next.js 14 (App Router), TypeScript, Tailwind CSS — RTL עברית מלא')
bullet('Backend: API Routes של Next.js (Serverless Functions)')
bullet('DB + Auth: Supabase (PostgreSQL + Row-Level Security + Google OAuth)')
bullet('תמלול: Deepgram Nova-2 עברית, fallback: Whisper Large')
bullet('AI: Anthropic Claude (claude-sonnet-4-6) — הפקת רשומות ומסמכים')
bullet('Hosting: Vercel (CI/CD אוטומטי מ-GitHub)')

heading('5.2 אבטחה', level=2)
bullet('Row-Level Security (RLS) — נתוני קליניקה נפרדים לחלוטין')
bullet('הזמנה-בלבד — אין הרשמה חופשית')
bullet('Admin Client עם Service Role Key לפעולות מנהל בלבד')
bullet('Env vars מאובטחים ב-Vercel (לא ב-repo)')

heading('5.3 בסיס הנתונים — טבלאות עיקריות', level=2)

tbl2 = doc.add_table(rows=8, cols=2)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['טבלה', 'תוכן עיקרי']):
    cell = tbl2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_run(cell.paragraphs[0].runs[0])

db_rows = [
    ['clinics',         'שם, slug, settings JSONB (template_id, לוגו)'],
    ['clinic_members',  'user_id, clinic_id, role, status, created_at'],
    ['profiles',        'full_name, phone, avatar_url'],
    ['patients',        'פרטים אישיים, קופה, אבחנה, סטטוס, bituach_leumi'],
    ['treatments',      'subjective/objective/assessment/plan, vas, note JSONB, template_id'],
    ['documents',       'type, title, content, status (draft/final)'],
    ['invitations',     'email, clinic_id, role, token, status'],
]
for r_idx, row_data in enumerate(db_rows):
    for c_idx, val in enumerate(row_data):
        cell = tbl2.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if cell.paragraphs[0].runs:
            set_rtl_run(cell.paragraphs[0].runs[0])

set_table_rtl(tbl2)
doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════
# 6. מצב נוכחי ומה בפיתוח
# ═══════════════════════════════════════════════════════════════════════════
heading('6. מצב נוכחי (יוני 2026)')

heading('6.1 מה עובד בפרודקשן', level=2)
bullet('✅ אימות, הזמנות, ניהול קליניקות ומשתמשים')
bullet('✅ ניהול מטופלים מלא')
bullet('✅ AI Scribe — הקלטה, תמלול עברית, יצירת רשומה מובנית')
bullet('✅ פקודות קוליות בעברית ("פרקסיס תרשום" / "פרקסיס סיים")')
bullet('✅ 15 תבניות קליניות מובנות לכל תחומי הפרא-רפואה + תבנית מותאמת')
bullet('✅ תבנית חלה על כל מוצר: טופס ידני, AI, היסטוריה, צ\'אט')
bullet('✅ מסמכים (טיוטה / סופי)')
bullet('✅ צ\'אט AI קליני')
bullet('✅ תרשים VAS (מגמת כאב) למטופל')
bullet('✅ Onboarding Center עם progress ring')
bullet('✅ Super Admin: ניהול קליניקות, מחיקה, מעבר בין קליניקות')
bullet('✅ קליניקת דמו עם נתונים מציאותיים (6 מטופלים, 14 טיפולים)')
bullet('✅ Landing page עם ROI calculator')

heading('6.2 מה בתכנון / חסר', level=2)
bullet('⬜ אנליטיקות מורחבות — לוח בקרה מנהלי, השוואות מטפלים, KPIs')
bullet('⬜ יצירת מסמכים ב-AI מלא (כרגע: עורך חינמי)')
bullet('⬜ תזמון תורים / יומן')
bullet('⬜ אפליקציית מובייל (iOS / Android)')
bullet('⬜ חיוב / Subscription (Stripe integration)')
bullet('⬜ חתימה דיגיטלית על מסמכים')
bullet('⬜ ייצוא PDF / Word לרשומות')
bullet('⬜ אינטגרציה עם מערכות HIS / EMR קיימות')
bullet('⬜ תיעוד מרחוק (Telehealth)')
bullet('⬜ תפיסת תמונות / X-ray annotation')
bullet('⬜ persistence של onboarding בין מכשירים (DB במקום localStorage)')

divider()

# ═══════════════════════════════════════════════════════════════════════════
# 7. זרימות משתמש
# ═══════════════════════════════════════════════════════════════════════════
heading('7. זרימות משתמש עיקריות')

heading('7.1 תיעוד טיפול — Scribe', level=2)
for s in [
    '1. כניסה לדף Scribe',
    '2. בחירת מטופל מהרשימה',
    '3. אמירת "פרקסיס תרשום" (או לחיצת כפתור)',
    '4. ניהול שיחת הטיפול בחדר',
    '5. אמירת "פרקסיס סיים"',
    '6. AI מתמלל ויוצר רשומה קלינית תוך ~15 שניות',
    '7. עיון, עריכה, הוספת VAS',
    '8. לחיצת "שמור רשומה"',
]:
    bullet(s)

heading('7.2 הוספת מטפל חדש', level=2)
for s in [
    '1. מנהל נכנס ל"משתמשים והרשאות"',
    '2. הזנת כתובת מייל + תפקיד',
    '3. מטפל מקבל מייל הזמנה',
    '4. הזמנה מקשרת לכניסה עם גוגל',
    '5. לאחר כניסה — חבר צוות פעיל',
]:
    bullet(s)

heading('7.3 שינוי תבנית תיעוד', level=2)
for s in [
    '1. מנהל נכנס להגדרות → תבנית תיעוד',
    '2. בחירת תחום מקצועי (פיזיותרפיה / קלינאות תקשורת / ריפוי בעיסוק וכו\')',
    '3. בחירת תבנית ספציפית מהקבוצה, או הדבקת דגימה לתבנית מותאמת',
    '4. לחיצת "בחר תבנית"',
    '5. כל התיעוד הבא (Scribe + טפסים ידניים) משתמש בתבנית החדשה',
]:
    bullet(s)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# 8. הגדרות API ואינטגרציות
# ═══════════════════════════════════════════════════════════════════════════
heading('8. אינטגרציות חיצוניות')

tbl3 = doc.add_table(rows=5, cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['שירות', 'מטרה', 'מפתח סביבה']):
    cell = tbl3.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_run(cell.paragraphs[0].runs[0])

int_rows = [
    ['Supabase',      'DB, Auth, Storage, RLS',     'NEXT_PUBLIC_SUPABASE_URL / ANON_KEY / SERVICE_ROLE_KEY'],
    ['Deepgram',      'תמלול קולי בעברית',           'DG_KEY'],
    ['Anthropic',     'יצירת רשומות ומסמכים',        'CL_KEY'],
    ['Google OAuth',  'כניסת משתמשים',               'מוגדר ב-Supabase Auth'],
]
for r_idx, row_data in enumerate(int_rows):
    for c_idx, val in enumerate(row_data):
        cell = tbl3.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if cell.paragraphs[0].runs:
            set_rtl_run(cell.paragraphs[0].runs[0])

set_table_rtl(tbl3)
doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════
# 9. שאלות פתוחות לדיון
# ═══════════════════════════════════════════════════════════════════════════
heading('9. שאלות פתוחות לדיון')

questions = [
    ('מודל עסקי', 'מנוי חודשי פר מטפל? פר קליניקה? Freemium? מה תמחור ראשוני?'),
    ('מובייל', 'האם ל-PWA מספיק? iOS/Android native? מתי?'),
    ('HIPAA / GDPR', 'מה רמת ציות הנדרשת? האם נדרש Data Processing Agreement עם Supabase/Anthropic/Deepgram?'),
    ('שפות נוספות', 'האם לתמוך בערבית? אנגלית? לאיזה שוק מכוונים בשנה 2?'),
    ('אינטגרציה עם EMR', 'כלומדית? מירב? Clalit HIS? מה הדחיפות?'),
    ('חתימה ואישור רשומה', 'האם נדרש workflow "רופא מאשר" לצורך ביטוח לאומי?'),
    ('Offline', 'האם נדרש תמיכה ללא אינטרנט (אזורים כפריים, מרפאות שדה)?'),
    ('Export / Backup', 'ייצוא נתונים לקליניקה עצמה? PDF אוטומטי?'),
    ('הרחבת תחומים', 'אופטומטריה, כירורגיה שיניים, עיסוי רפואי — מתי ואיך?'),
]
for title, desc in questions:
    p = doc.add_paragraph()
    set_rtl_para(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f'❓ {title}: ')
    r1.bold = True
    r1.font.size = Pt(11)
    set_rtl_run(r1)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)
    set_rtl_run(r2)

divider()

# ═══════════════════════════════════════════════════════════════════════════
# 10. מדדי הצלחה
# ═══════════════════════════════════════════════════════════════════════════
heading('10. מדדי הצלחה (KPIs)')

bullet('זמן תיעוד ממוצע לטיפול < 2 דקות (כולל עריכה)')
bullet('שביעות רצון מטפלים ≥ 4.5/5')
bullet('דיוק תמלול עברית > 90%')
bullet('אחוז ניסויי Scribe שהסתיימו בשמירה > 80%')
bullet('Churn חודשי < 5%')
bullet('זמן קליטה (Onboarding) ≤ 10 דקות לקליניקה חדשה')

divider()

# ── Footer note ────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('praxisAI · מסמך פנימי לשיתוף צוות בלבד · יוני 2026')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(148, 163, 184)
r.italic = True
set_rtl_run(r)

# ── Save ───────────────────────────────────────────────────────────────────
out = '/home/user/praxisai/praxisai-spec.docx'
doc.save(out)
print(f'Saved: {out}')
